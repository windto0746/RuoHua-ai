# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_chinese_font(run, font_name='微软雅黑', font_size=11):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading('', level=level)
    run = heading.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0x8A, 0x9B)  # 珊瑚粉
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x9B, 0x6B, 0x9E)  # 紫色
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)  # 深灰
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading

def add_paragraph(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    
    # 表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = '微软雅黑'
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # 表头背景色
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'FF8A9B')
        header_cells[i]._tc.get_or_add_tcPr().append(shading)
    
    # 数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
        # 交替背景色
        if row_idx % 2 == 1:
            for cell in row_cells:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'FFF5F7')
                cell._tc.get_or_add_tcPr().append(shading)
    
    return table

# 创建文档
doc = Document()

# 标题
title = doc.add_heading('', 0)
run = title.add_run('「若花」AI文案生成器')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x8A, 0x9B)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

subtitle = doc.add_paragraph()
run = subtitle.add_run('——项目升级历程与产品经理素养分析')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x8E, 0x8E, 0x93)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 一、项目概述
add_heading(doc, '一、项目概述', 1)
add_paragraph(doc, '「若花」是一款面向小红书内容创作者的AI文案生成工具，用户只需输入主题关键词，即可快速生成适配平台的优质文案。')
doc.add_paragraph('项目从2026年4月9日启动，历经多个版本的迭代优化，最终形成了一套完整的「场景×风格」产品矩阵。')
doc.add_paragraph('当前版本已部署上线，访问地址：https://67fe5cm79wb4.space.minimaxi.com')

# 二、版本升级历程
add_heading(doc, '二、版本升级历程（箭头式时间线）', 1)

# v1.0
add_heading(doc, '▼ v1.0 基础版本（4月9日）', 3)
add_paragraph(doc, '┌ 核心功能：', indent=True)
add_paragraph(doc, '• 单一场景生成\n• 固定文案模板\n• 基础输入输出', indent=True)
add_paragraph(doc, '└ 技术实现：Python脚本 + DashScope API调用', indent=True)

# v2.0
add_heading(doc, '▶ v2.0 场景扩展（4月11日）', 3)
add_paragraph(doc, '┌ 核心升级：', indent=True)
add_paragraph(doc, '• 新增「重叙述型」场景（80-150字）\n• 新增「重总结型」场景（10-20字）\n• 设计四风格体系（感情系/抽象系/启发系/种草系）', indent=True)
add_paragraph(doc, '└ 用户需求：不同视频类型需要不同的文案风格\n└ 解决方案：「场景×风格」矩阵架构', indent=True)

# v3.0
add_heading(doc, '▶ v3.0 交互优化（4月11日）', 3)
add_paragraph(doc, '┌ 核心升级：', indent=True)
add_paragraph(doc, '• 新增多轮对话机制\n• 支持用户反馈修改\n• 实现对话历史记录', indent=True)
add_paragraph(doc, '└ 用户需求：AI初稿难以完全符合预期，需要迭代优化\n└ 解决方案：生成→反馈→优化→...→保存的交互流程', indent=True)

# v4.0
add_heading(doc, '▶ v4.0 长篇场景（4月12日）', 3)
add_paragraph(doc, '┌ 核心升级：', indent=True)
add_paragraph(doc, '• 新增「长篇故事型」场景（500-2000字）\n• 完善Prompt模板体系\n• 增加敏感词过滤', indent=True)
add_paragraph(doc, '└ 用户需求：部分视频以文案为核心内容，需要完整故事\n└ 解决方案：三大场景全面覆盖不同内容需求', indent=True)

# v5.0
add_heading(doc, '▶ v5.0 Gradio Web版（4月12日）', 3)
add_paragraph(doc, '┌ 核心升级：', indent=True)
add_paragraph(doc, '• 基于Gradio构建Web界面\n• 热门话题推荐\n• 历史记录存储', indent=True)
add_paragraph(doc, '└ 用户需求：从命令行走向可视化操作\n└ 解决方案：现代化Web界面，降低使用门槛', indent=True)

# v6.0
add_heading(doc, '▶ v6.0 「若花」界面重塑（4月14日）', 3)
add_paragraph(doc, '┌ 核心升级：', indent=True)
add_paragraph(doc, '• 参考豆包设计风格\n• 樱花主题图标\n• 聊天式交互界面\n• 消息气泡展示', indent=True)
add_paragraph(doc, '└ 用户需求：追求更友好的用户体验和视觉美感\n└ 解决方案：定制化UI设计，融入品牌元素', indent=True)

# v7.0
add_heading(doc, '▶ v7.0 对话系统重构（4月15日）', 3)
add_paragraph(doc, '┌ 核心升级：', indent=True)
add_paragraph(doc, '• 新建对话功能\n• 版本保留机制（修改/重新生成追加新版本）\n• 对话历史管理\n• 点击跳转对话', indent=True)
add_paragraph(doc, '└ 用户需求：每次修改不覆盖历史，希望保留所有版本\n└ 解决方案：卡片式对话设计，版本标签体系', indent=True)

# 三、用户需求与解决方案对照表
add_heading(doc, '三、用户需求与解决方案对照表', 1)
headers = ['序号', '用户原始需求', '解决方案', '体现的产品思维']
rows = [
    ['1', '不同视频类型需要不同文案', '设计「场景×风格」矩阵架构', '场景细分能力'],
    ['2', 'AI初稿不够完美', '多轮对话交互机制', '迭代思维'],
    ['3', '希望生成完整故事内容', '新增长篇故事型场景', '需求挖掘能力'],
    ['4', '命令行使用不方便', 'Gradio Web界面', '用户体验意识'],
    ['5', '界面不够美观', '豆包风格+樱花主题', '审美与设计能力'],
    ['6', '修改后旧版本消失', '版本保留+卡片式对话', '数据管理思维'],
    ['7', '历史记录查找不便', '点击跳转+对话为单位', '信息架构能力'],
    ['8', '同一主题生成重复', '唯一编号+随机提示', '问题解决能力'],
]
add_table(doc, headers, rows)

# 四、产品经理核心素养分析
add_heading(doc, '四、产品经理核心素养分析', 1)
add_paragraph(doc, '通过「若花」项目的完整迭代过程，充分体现了以下AI产品经理的核心素养：')

# 8项素养
skills = [
    ('① 场景洞察能力', '从用户实际使用场景出发，将单一的「生成文案」需求，细分为「重叙述」「重总结」「长篇故事」三大场景。简历可写：「善于从用户实际使用场景出发进行产品功能拆解」'),
    ('② 需求优先级判断', '优先实现核心生成功能，再逐步完善交互细节。简历可写：「具备产品迭代优先级判断能力，擅于在资源有限的情况下聚焦核心功能」'),
    ('③ 迭代优化思维', '从「一次性生成」到「多轮交互优化」，体现了对AI产品本质的理解。简历可写：「理解AI产品的不确定性，擅于设计人机协作的迭代优化机制」'),
    ('④ 用户体验意识', '从命令行到Web界面，从基础界面到豆包风格，每一次升级都体现了对用户体验的持续关注。简历可写：「具备用户体验优先的产品思维」'),
    ('⑤ Prompt工程能力', '设计了完整的Prompt模板体系，每个场景有明确的字数、结构、格式要求。简历可写：「具备Prompt工程实战经验，能够针对不同业务场景设计高效AI提示词」'),
    ('⑥ 问题解决能力', '面对「生成重复」和「版本覆盖」等问题，提出了创新的解决方案。简历可写：「擅于发现并解决产品问题，具备从问题到方案的闭环思维」'),
    ('⑦ 审美与设计能力', '从零开始设计「若花」的视觉形象，选用樱花作为品牌元素。简历可写：「具备基础的产品设计能力，能够从0到1规划产品的视觉和交互风格」'),
    ('⑧ 数据驱动思维', '通过本地存储记录用户对话历史，支持对话管理和历史回溯。简历可写：「具备数据意识，关注用户行为数据的采集与应用」'),
]

for title, desc in skills:
    add_heading(doc, title, 3)
    add_paragraph(doc, desc, indent=True)

# 五、面试经典问答准备
add_heading(doc, '五、面试经典问答准备', 1)

add_heading(doc, 'Q：介绍一下你在「若花」项目中的角色和贡献？', 3)
add_paragraph(doc, '「若花」是我独立完成产品设计和开发的项目。核心贡献包括：提出「场景×风格」的产品矩阵架构、设计多轮交互机制、从0到1规划产品的视觉风格、实现版本管理系统。')

add_heading(doc, 'Q：这个项目最大的挑战是什么？你是如何解决的？', 3)
add_paragraph(doc, '最大的挑战是让AI生成的文案真正「能用」。后来我意识到，AI产品的本质不是一次完美，而是「生成-反馈-优化」的迭代过程。所以我设计了多轮交互机制，用户可以不断提意见，AI基于对话历史优化结果。')

add_heading(doc, 'Q：你如何理解AI产品经理这个岗位？', 3)
add_paragraph(doc, '我认为AI产品经理的核心是「让AI能力真正解决用户问题」。有三个关键能力：①场景洞察②Prompt工程③迭代思维。「若花」项目就是很好的例子——从v1.0到v7.0，每个版本都是基于用户反馈和自我迭代产出的。')

# 六、技术栈
add_heading(doc, '六、项目技术栈总结', 1)
headers = ['层级', '技术/工具', '说明']
rows = [
    ['编程语言', 'Python', '基础开发能力'],
    ['AI模型', '通义千问 (DashScope)', '大语言模型API调用'],
    ['前端框架', 'HTML/CSS/JavaScript', '响应式Web界面'],
    ['UI组件', 'Gradio', '快速原型开发'],
    ['版本管理', 'Git', '代码版本控制'],
    ['部署平台', 'Minimax Spaces', '在线访问'],
]
add_table(doc, headers, rows)

# 保存文档
output_path = r'C:\Users\zjw\.minimax-agent-cn\projects\5\若花项目升级历程.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
