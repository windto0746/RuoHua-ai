"""
创建简历项目经历Word文档
若花文案生成器 - 产品迭代历程
使用箭头式排版体现版本演进
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 创建文档
doc = Document()

# 设置文档默认字体
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 设置页边距
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# ========== 标题部分 ==========
# 主标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('若花文案生成器')
run.font.name = 'Microsoft YaHei'
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(44, 62, 80)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 副标题
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('产品迭代历程')
run.font.name = 'Microsoft YaHei'
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(108, 114, 127)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 描述性副标题
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('AI驱动的文案生成工具，从命令行脚本演进为完整产品')
run.font.name = 'Microsoft YaHei'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)
run.font.italic = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 添加段落间距
doc.add_paragraph()

# ========== 版本演进部分 ==========
# 定义版本数据
versions = [
    {
        'version': 'v1.0',
        'name': '基础版本',
        'upgrade': '命令行脚本调用API生成基础文案',
        'skill': '需求洞察 —— 从用户"想要快速生成文案"的核心诉求出发'
    },
    {
        'version': 'v2.0',
        'name': '场景扩展',
        'upgrade': '新增「重叙述型」「重总结型」场景，设计四风格体系',
        'skill': '场景细分 —— 将单一需求拆解为「场景×风格」矩阵'
    },
    {
        'version': 'v3.0',
        'name': '交互优化',
        'upgrade': '多轮对话机制，支持用户反馈修改',
        'skill': '迭代思维 —— 理解AI产品需要"生成-反馈-优化"循环'
    },
    {
        'version': 'v4.0',
        'name': '长篇场景',
        'upgrade': '新增「长篇故事型」场景，完善Prompt模板',
        'skill': '需求挖掘 —— 挖掘用户对完整故事内容的需求'
    },
    {
        'version': 'v5.0',
        'name': 'Web界面',
        'upgrade': 'Gradio Web界面，热门话题推荐',
        'skill': '用户体验 —— 从命令行走向可视化，降低使用门槛'
    },
    {
        'version': 'v6.0',
        'name': '界面重塑',
        'upgrade': '豆包风格+樱花主题，聊天式交互',
        'skill': '审美设计 —— 从0到1规划产品视觉风格，品牌元素融入'
    },
    {
        'version': 'v7.0',
        'name': '对话重构',
        'upgrade': '版本保留机制，新建对话，历史管理',
        'skill': '数据管理 —— 解决版本覆盖问题，信息架构能力'
    }
]


def add_arrow_line(doc):
    """添加箭头分隔线"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 使用Unicode箭头字符
    run = p.add_run('─' * 25 + '►' + '─' * 25)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(180, 180, 180)


# 遍历每个版本
for i, v in enumerate(versions):
    # 添加版本卡片背景（使用段落边框模拟）
    if i > 0:
        add_arrow_line(doc)
    
    # 版本号和名称
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 版本号
    run = version_para.add_run(f"{v['version']} {v['name']}")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(52, 152, 219)  # 蓝色
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 升级方面（上方）
    upgrade_para = doc.add_paragraph()
    upgrade_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 标签
    run = upgrade_para.add_run('「升级方面」')
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(149, 165, 166)
    run.font.italic = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    upgrade_para.add_run('\n')
    
    # 升级内容
    run = upgrade_para.add_run(v['upgrade'])
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(44, 62, 80)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 素养体现（下方）
    skill_para = doc.add_paragraph()
    skill_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 标签
    run = skill_para.add_run('「体现素养」')
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(149, 165, 166)
    run.font.italic = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    skill_para.add_run('\n')
    
    # 素养内容 - 使用强调色
    run = skill_para.add_run(v['skill'])
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(39, 174, 96)  # 绿色
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 添加小间距
    doc.add_paragraph()


# ========== 底部总结 ==========
doc.add_paragraph()

summary = doc.add_paragraph()
summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = summary.add_run('7次迭代，完整经历产品从0到1的全过程')
run.font.name = 'Microsoft YaHei'
run.font.size = Pt(11)
run.font.bold = True
run.font.color.rgb = RGBColor(44, 62, 80)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 保存文档
output_path = r'C:\Users\zjw\.minimax-agent-cn\projects\5\若花项目经历-简历版.docx'
doc.save(output_path)

print(f"文档已创建: {output_path}")